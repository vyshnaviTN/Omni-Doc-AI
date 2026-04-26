import os
import re
import unicodedata
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pytesseract
from PIL import Image
import base64
from io import BytesIO
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage

# Tesseract OCR Configuration for Windows
DEFAULT_TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(DEFAULT_TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = DEFAULT_TESSERACT_PATH

# ---------------------------------------------------------------------------
# Chunk Quality Constants
# ---------------------------------------------------------------------------

MIN_CHUNK_LENGTH = 30           # Discard chunks shorter than this
MAX_SYMBOL_RATIO = 0.40         # Discard if >40% non-alphanumeric (OCR noise)
CHUNK_SIZE = 800                # ~400 tokens for precise retrieval
CHUNK_OVERLAP = 160             # ~80 tokens overlap for continuity


# ---------------------------------------------------------------------------
# Text Cleaning
# ---------------------------------------------------------------------------

def _is_mostly_ascii(text: str, min_ratio: float = 0.80) -> bool:
    """Return True if at least `min_ratio` of printable chars are ASCII."""
    printable = [c for c in text if not unicodedata.category(c).startswith('C')]
    if not printable:
        return False
    ascii_count = sum(1 for c in printable if ord(c) < 128)
    return (ascii_count / len(printable)) >= min_ratio


def _clean_text(text: str) -> str:
    """
    Multi-stage text cleaner:
    1. Remove null bytes and control characters.
    2. Strip PDF/Unicode garbage (Private Use Area, zero-width, BOM).
    3. Collapse repeated non-alphanumeric symbols.
    4. Normalize whitespace.
    5. Discard if mostly non-ASCII junk.
    """
    if not text:
        return ""

    # Remove null bytes and control characters (keep newline/tab)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

    # Remove common PDF encoding garbage patterns
    text = re.sub(r'[\uf000-\uf8ff]', '', text)   # Private Use Area
    text = re.sub(r'[\u200b-\u200f\u202a-\u202e\ufeff]', '', text)  # Zero-width / BOM

    # Collapse repeated non-alphanumeric chars (e.g. "----", "####")
    text = re.sub(r'([^\w\s])\1{3,}', r'\1', text)

    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()

    # Final garbage check — if less than 75% of chars are ASCII, discard
    if not _is_mostly_ascii(text, min_ratio=0.75):
        return ""

    return text


def _is_noisy_chunk(text: str) -> bool:
    """
    Detect OCR noise / garbage chunks:
    - Too short to be meaningful.
    - High ratio of non-alphanumeric characters (garbled OCR output).
    """
    stripped = text.strip()

    if len(stripped) < MIN_CHUNK_LENGTH:
        return True

    # Count alphanumeric vs total printable characters
    alphanumeric = sum(1 for c in stripped if c.isalnum() or c.isspace())
    if len(stripped) > 0 and (alphanumeric / len(stripped)) < (1 - MAX_SYMBOL_RATIO):
        return True

    return False


# ---------------------------------------------------------------------------
# OCR Verification Storage
# ---------------------------------------------------------------------------

def _save_ocr_verification(img: Image.Image, text: str, source_filename: str, page: int = 1):
    """Save the OCR image and extracted text for development verification."""
    try:
        import time
        safe_filename = re.sub(r'[^a-zA-Z0-9_\-]', '_', source_filename)
        base_dir = os.path.dirname(os.path.dirname(__file__))
        verify_dir = os.path.join(base_dir, "ocr_verification")
        os.makedirs(verify_dir, exist_ok=True)
        
        timestamp = int(time.time())
        base_name = f"{timestamp}_{safe_filename}_page{page}"
        
        img_path = os.path.join(verify_dir, f"{base_name}.png")
        text_path = os.path.join(verify_dir, f"{base_name}.txt")
        
        # Save image
        img.save(img_path, format="PNG")
        
        # Save text
        with open(text_path, "w", encoding="utf-8") as f:
            f.write(text)
            
        print(f"[VERIFY] OCR saved to {verify_dir} for {base_name}")
    except Exception as e:
        print(f"[VERIFY] Failed to save OCR verification for {source_filename}: {e}")

# ---------------------------------------------------------------------------
# VLM Extraction Engine (Llama 3.2 Vision)
# ---------------------------------------------------------------------------

def _extract_text_via_vlm(img: Image.Image, prompt_type: str = "document") -> str:
    """
    Extract text using Llama 3.2 Vision via Ollama.
    Handles difficult handwriting and poor quality scans better than Tesseract.
    """
    try:
        # 1. Convert to Base64
        buffered = BytesIO()
        # Save as JPEG with high quality to reduce payload size while keeping detail
        img.save(buffered, format="JPEG", quality=90)
        img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")

        # 2. Initialize VLM (moondream)
        vision_llm = ChatOllama(model="moondream", temperature=0.0)

        # 3. Construct Prompt (Simplified for moondream)
        prompt_text = "Read all the text in this image and transcribe it exactly."
        
        # 4. Process in slices to prevent moondream from failing on dense pages
        w, h = img.size
        # If the image is tall (like a full page), split into 3 horizontal slices
        num_slices = 3 if h > w else 1
        slice_height = h // num_slices
        
        extracted_text = []
        print(f"[VLM] Sending image to moondream in {num_slices} slice(s)...")
        
        import time
        start = time.time()
        
        for i in range(num_slices):
            top = i * slice_height
            bottom = h if i == num_slices - 1 else (i + 1) * slice_height
            patch = img.crop((0, top, w, bottom))
            
            buffered = BytesIO()
            patch.save(buffered, format="JPEG", quality=85)
            img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
            
            message = HumanMessage(
                content=[
                    {"type": "text", "text": prompt_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_base64}"}},
                ]
            )
            
            try:
                res = vision_llm.invoke([message])
                text = res.content.strip()
                if text:
                    extracted_text.append(text)
            except Exception as e:
                print(f"[VLM] Error on slice {i+1}: {e}")
                
        elapsed = time.time() - start
        final_text = "\n\n".join(extracted_text).strip()
        print(f"[VLM] Extraction complete in {elapsed:.2f}s. Extracted {len(final_text)} chars.")
        
        return final_text
    except Exception as e:
        print(f"[VLM] Error during vision extraction: {e}")
        return ""

# ---------------------------------------------------------------------------
# PDF Extraction — PyMuPDF (Primary) with OCR fallback
# ---------------------------------------------------------------------------

def _extract_pdf_fitz(file_path: str, filename: str) -> list:
    """
    Primary PDF extractor using PyMuPDF (fitz).
    Produces reliable UTF-8 text without the encoding issues of PyPDF/pypdf.
    Falls back to OCR per-page if extracted text is too short.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    docs = []
    needs_ocr_pages = []

    for i in range(len(doc)):
        page = doc.load_page(i)
        text = page.get_text("text")
        cleaned = _clean_text(text)

        if len(cleaned) > MIN_CHUNK_LENGTH:
            docs.append(Document(
                page_content=cleaned,
                metadata={"page": i + 1, "source": filename}
            ))
        else:
            # This page had no selectable text — queue for OCR
            needs_ocr_pages.append((i, page))

    # Run OCR on pages that had no text
    if needs_ocr_pages:
        print(f"[{filename}] {len(needs_ocr_pages)} page(s) need OCR fallback.")
        for i, page in needs_ocr_pages:
            try:
                import fitz
                # High DPI (300) required for reliability
                pix = page.get_pixmap(dpi=300, colorspace=fitz.csRGB, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Try VLM first for handwriting/scanned pages
                ocr_text = _extract_text_via_vlm(img)
                extraction_method = "vlm"

                # Fallback to Tesseract if VLM fails or returns nothing
                if not ocr_text or len(ocr_text.strip()) < 10:
                    print(f"[{filename}] VLM failed or too short, falling back to Tesseract on page {i + 1}...")
                    # Pre-process for Tesseract
                    gray_img = img.convert('L')
                    gray_img = gray_img.point(lambda p: 255 if p > 180 else 0)
                    ocr_text = pytesseract.image_to_string(gray_img, lang='eng', config='--psm 3')
                    extraction_method = "ocr-tesseract"
                
                # Save verification copy
                _save_ocr_verification(img, ocr_text, filename, page=i + 1)
                
                cleaned = _clean_text(ocr_text)
                if cleaned and len(cleaned) > MIN_CHUNK_LENGTH:
                    docs.append(Document(
                        page_content=cleaned,
                        metadata={
                            "page": i + 1, 
                            "source": filename, 
                            "extraction": extraction_method,
                            "vlm_processed": (extraction_method == "vlm")
                        }
                    ))
                else:
                    print(f"[{filename}] OCR page {i + 1} discarded (too short/noisy).")
            except Exception as ocr_err:
                print(f"[{filename}] OCR failed on page {i + 1}: {ocr_err}")

    doc.close()
    return docs


def _extract_pdf_ocr_full(file_path: str, filename: str) -> list:
    """
    Full OCR fallback using pypdfium2 + pytesseract.
    Used when PyMuPDF cannot open or process the PDF at all (e.g., corrupted files).
    Converts each page to an image and applies OCR without requiring external Poppler binaries.
    """
    try:
        import pypdfium2 as pdfium
    except ImportError:
        print(f"[{filename}] pypdfium2 not installed — OCR fallback unavailable.")
        return []

    docs = []
    try:
        pdf = pdfium.PdfDocument(file_path)
    except Exception as e:
        print(f"[{filename}] pypdfium2 loading failed: {e}")
        raise ValueError(f"Unable to extract readable text from '{filename}'")

    for i in range(len(pdf)):
        try:
            page = pdf[i]
            bitmap = page.render(scale=2.77) # ~200 DPI (72 * 2.77)
            img = bitmap.to_pil()
            
            # Try VLM first
            ocr_text = _extract_text_via_vlm(img)
            extraction_method = "vlm"
            
            if not ocr_text or len(ocr_text.strip()) < 10:
                print(f"[{filename}] VLM failed/too short, falling back to Tesseract on full-page {i + 1}...")
                ocr_text = pytesseract.image_to_string(img, lang='eng')
                extraction_method = "ocr-tesseract"
            
            # Save verification copy
            _save_ocr_verification(img, ocr_text, filename, page=i + 1)
            
            cleaned = _clean_text(ocr_text)
            if cleaned and len(cleaned) > MIN_CHUNK_LENGTH:
                docs.append(Document(
                    page_content=cleaned,
                    metadata={
                        "page": i + 1, 
                        "source": filename, 
                        "extraction": extraction_method,
                        "vlm_processed": (extraction_method == "vlm")
                    }
                ))
            else:
                print(f"[{filename}] OCR full-page {i + 1} discarded (too short/noisy).")
        except Exception as ocr_err:
            print(f"[{filename}] OCR failed on full-page {i + 1}: {ocr_err}")

    return docs



# ---------------------------------------------------------------------------
# DOCX Extraction
# ---------------------------------------------------------------------------

def _load_docx(file_path: str, filename: str) -> list:
    """Parse DOCX using python-docx, preserving paragraph structure."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        cleaned = _clean_text(full_text)
        if not cleaned:
            raise ValueError("No usable text found in the DOCX document.")
        return [Document(page_content=cleaned, metadata={"page": 1, "source": filename})]
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")


# ---------------------------------------------------------------------------
# Image OCR
# ---------------------------------------------------------------------------

def _extract_image(file_path: str, filename: str) -> list:
    """Extract text from image files via Tesseract OCR."""
    try:
        img = Image.open(file_path)
        
        # Handle alpha channel
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
            
        # Try VLM first
        ocr_text = _extract_text_via_vlm(img)
        extraction_method = "vlm"
        
        if not ocr_text or len(ocr_text.strip()) < 10:
            print(f"[{filename}] VLM failed/too short, falling back to Tesseract for image...")
            ocr_text = pytesseract.image_to_string(img, lang='eng')
            extraction_method = "ocr-tesseract"
        
        # Save verification copy
        _save_ocr_verification(img, ocr_text, filename, page=1)
        
        cleaned = _clean_text(ocr_text)
        if not cleaned:
            raise ValueError("No readable text found in the image.")
        return [Document(
            page_content=cleaned, 
            metadata={
                "page": 1, 
                "source": filename, 
                "extraction": extraction_method,
                "vlm_processed": (extraction_method == "vlm")
            }
        )]
    except Exception as e:
        print(f"[ERROR] Image OCR failed: {e}")
        raise ValueError(f"Image OCR failed: {e}")


# ---------------------------------------------------------------------------
# Main Entry Point
# ---------------------------------------------------------------------------

def process_file(file_path: str, filename: str, extension: str) -> list:
    """
    Loads a file, cleans text, splits into semantic chunks, and returns them.
    Supported: .pdf, .docx, .txt, .md, .png, .jpg, .jpeg

    Uses PyMuPDF as primary PDF extractor with per-page OCR fallback.
    If PyMuPDF fails entirely, falls back to pdf2image + pytesseract for
    full-document OCR (handles scanned/handwritten PDFs).
    """
    documents = []

    try:
        if extension == '.pdf':
            try:
                documents = _extract_pdf_fitz(file_path, filename)
                if not documents:
                    raise ValueError("No readable chunks generated.")
            except Exception as fitz_err:
                print(f"[{filename}] PyMuPDF extraction failed ({fitz_err}), attempting full OCR fallback...")
                documents = _extract_pdf_ocr_full(file_path, filename)

            if not documents:
                raise ValueError(
                    "Unable to extract readable text from this PDF. "
                    "Ensure Tesseract OCR is installed and the document is readable."
                )

        elif extension == '.docx':
            documents = _load_docx(file_path, filename)

        elif extension in ('.txt', '.md'):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    raw = f.read()
            cleaned = _clean_text(raw)
            if not cleaned:
                raise ValueError("No usable text found in this file.")
            documents = [Document(page_content=cleaned, metadata={"page": 1, "source": filename})]

        elif extension in ('.png', '.jpg', '.jpeg'):
            documents = _extract_image(file_path, filename)

        else:
            raise ValueError(
                f"Unsupported file type: '{extension}'. "
                "Supported formats: PDF, DOCX, TXT, MD, PNG, JPG, JPEG."
            )

        # Ensure all docs have the correct source metadata
        for doc in documents:
            doc.metadata["source"] = filename
            doc.metadata.setdefault("extraction", "text")

        # Remove any empty or near-empty documents that slipped through
        documents = [d for d in documents if len(d.page_content.strip()) > MIN_CHUNK_LENGTH]

        if not documents:
            raise ValueError(
                "Document processed but no usable text was found after cleaning. "
                "The file may be blank, image-only, or use an unsupported encoding."
            )

        # Semantic chunking — ~400 tokens with ~80 token overlap for precise retrieval.
        # Sentence-boundary separators prevent breaking mid-sentence.
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""],
        )
        chunks = splitter.split_documents(documents)

        # Post-split quality filter: remove noisy/garbage chunks
        clean_chunks = []
        discarded = 0
        for chunk in chunks:
            if _is_noisy_chunk(chunk.page_content):
                discarded += 1
                continue
            clean_chunks.append(chunk)

        if discarded > 0:
            print(f"[{filename}] Discarded {discarded} noisy/short chunks.")

        # Stamp chunk_index so citations can reference precise locations
        for i, chunk in enumerate(clean_chunks):
            chunk.metadata["chunk_index"] = i

        print(f"[OK] '{filename}': {len(clean_chunks)} clean chunks extracted.")
        return clean_chunks

    except Exception as e:
        print(f"[ERROR] processing '{filename}': {e}")
        raise
